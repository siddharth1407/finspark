"""
Document parsing service - handles PDF and text extraction
"""
import io
import os
import logging
from typing import Tuple, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parse various document formats and extract text.
    Supports: PDF, TXT, DOCX
    """

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "text/plain": "txt",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    }

    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def parse(
        self,
        file_content: bytes,
        filename: str,
        content_type: str
    ) -> Tuple[str, dict]:
        """
        Parse document and extract text.

        Returns:
            Tuple of (extracted_text, metadata)
        """
        file_type = self.SUPPORTED_TYPES.get(content_type, "unknown")

        if file_type == "pdf":
            return await self._parse_pdf(file_content, filename)
        elif file_type == "txt":
            return await self._parse_text(file_content, filename)
        elif file_type == "docx":
            return await self._parse_docx(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {content_type}")

    async def _parse_pdf(
        self,
        content: bytes,
        filename: str
    ) -> Tuple[str, dict]:
        """Extract text from PDF."""
        try:
            import pdfplumber

            text_parts = []
            metadata = {
                "filename": filename,
                "type": "pdf",
                "pages": 0
            }

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                metadata["pages"] = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)
            metadata["char_count"] = len(full_text)
            metadata["word_count"] = len(full_text.split())

            logger.info(
                f"Extracted {metadata['word_count']} words from {filename}")
            return full_text, metadata

        except ImportError:
            logger.warning("pdfplumber not installed, trying PyPDF2")
            return await self._parse_pdf_fallback(content, filename)
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise

    async def _parse_pdf_fallback(
        self,
        content: bytes,
        filename: str
    ) -> Tuple[str, dict]:
        """Fallback PDF parsing with PyPDF2."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(content))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            full_text = "\n\n".join(text_parts)
            return full_text, {
                "filename": filename,
                "type": "pdf",
                "pages": len(reader.pages),
                "char_count": len(full_text)
            }
        except Exception as e:
            logger.error(f"PDF fallback parsing error: {e}")
            raise

    async def _parse_text(
        self,
        content: bytes,
        filename: str
    ) -> Tuple[str, dict]:
        """Parse plain text file."""
        # Try different encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                return text, {
                    "filename": filename,
                    "type": "txt",
                    "encoding": encoding,
                    "char_count": len(text),
                    "word_count": len(text.split())
                }
            except UnicodeDecodeError:
                continue

        raise ValueError("Could not decode text file")

    async def _parse_docx(
        self,
        content: bytes,
        filename: str
    ) -> Tuple[str, dict]:
        """Parse DOCX file."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also get text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)
            return full_text, {
                "filename": filename,
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "char_count": len(full_text)
            }
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise

    def save_file(self, content: bytes, filename: str) -> str:
        """Save uploaded file to disk."""
        import hashlib

        # Create unique filename
        file_hash = hashlib.md5(content).hexdigest()[:8]
        safe_name = f"{file_hash}_{filename}"
        file_path = self.upload_dir / safe_name

        with open(file_path, "wb") as f:
            f.write(content)

        return str(file_path)


# Singleton
_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
